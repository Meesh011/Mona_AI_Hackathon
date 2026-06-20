import os
import sys
import pdfplumber
import pandas as pd
import pytesseract
from PIL import Image
from docx import Document

# ── FIX 1: Hard-coded Windows path replaced with auto-detection ──────────────
# On Windows, set this to r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# On Linux/Mac, tesseract is on PATH so no override needed.
if sys.platform == "win32":
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# else: uses system tesseract automatically


def read_pdf(path):
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text


def read_docx(path):
    doc = Document(path)

    # ── FIX 2: Original code only read paragraphs, missing ALL table content ──
    # Vendor name, invoice type, line items etc. live in tables in these files,
    # so the classifier had almost no signal to work with for DOCX invoices.
    parts = []

    # Paragraphs (headers, addresses outside tables)
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Tables (where the real invoice data lives)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n".join(parts)


def read_image(path):
    image = Image.open(path)
    return pytesseract.image_to_string(image, lang="eng+deu")


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()

    # ── FIX 3: Skip non-invoice files like the manifest CSV ──────────────────
    if ext == ".csv":
        return ""

    if ext == ".pdf":
        return read_pdf(path)

    elif ext == ".docx":
        return read_docx(path)

    elif ext in [".png", ".jpg", ".jpeg"]:
        return read_image(path)

    elif ext in [".xlsx", ".xls"]:
        sheets = pd.read_excel(path, sheet_name=None)
        text = ""
        for _, df in sheets.items():
            text += df.to_string()
        return text

    return ""
